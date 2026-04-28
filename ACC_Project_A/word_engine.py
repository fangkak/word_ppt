from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


def md_to_word(md_file, word_file):
    """
    将 Markdown 文件转换为 Word 文档
    
    Args:
        md_file: Markdown 源文件路径
        word_file: Word 输出文件路径
    """
    # 读取 Markdown 文件
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置文档样式
    setup_styles(doc)
    
    # 解析 Markdown 并添加到 Word
    parse_markdown_to_word(doc, content)
    
    # 保存 Word 文件
    doc.save(word_file)
    print(f"✓ Word 文档已生成: {word_file}")


def setup_styles(doc):
    """
    设置 Word 文档样式
    """
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)


def parse_markdown_to_word(doc, content):
    """
    解析 Markdown 内容并添加到 Word 文档
    
    Args:
        doc: Document 对象
        content: Markdown 内容字符串
    """
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # 跳过空行
        if not line.strip():
            i += 1
            continue
        
        # 一级标题
        if line.startswith('# '):
            title = line.replace('# ', '').strip()
            heading = doc.add_heading(title, level=1)
            set_heading_color(heading, RGBColor(0, 51, 102))
        
        # 二级标题
        elif line.startswith('## '):
            title = line.replace('## ', '').strip()
            heading = doc.add_heading(title, level=2)
            set_heading_color(heading, RGBColor(0, 102, 204))
        
        # 三级标题
        elif line.startswith('### '):
            title = line.replace('### ', '').strip()
            heading = doc.add_heading(title, level=3)
            set_heading_color(heading, RGBColor(51, 102, 153))
        
        # 水平线
        elif line.strip() == '---':
            add_horizontal_line(doc)
        
        # 表格
        elif line.strip().startswith('|'):
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            i -= 1  # 回退一行，因为外层循环会再加1
            
            # 解析表格
            if len(table_lines) >= 3:  # 至少要有表头、分隔符、一行数据
                add_table_from_markdown(doc, table_lines)
        
        # 图片
        elif line.strip().startswith('!['):
            img_path = extract_image_path(line)
            if img_path:
                try:
                    doc.add_picture(img_path, width=Inches(6))
                    # 添加图片下方的空行
                    doc.add_paragraph()
                except Exception as e:
                    print(f"⚠️  添加图片失败 {img_path}: {e}")
                    doc.add_paragraph(f"[图片错误: {img_path}]")
        
        # 列表项
        elif line.strip().startswith('- '):
            # 收集列表项
            list_items = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                item = lines[i].strip().replace('- ', '', 1).strip()
                list_items.append(item)
                i += 1
            i -= 1
            
            # 添加列表
            for item in list_items:
                doc.add_paragraph(item, style='List Bullet')
        
        # 代码块
        elif line.strip().startswith('```'):
            # 收集代码块
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            # 添加代码块
            if code_lines:
                p = doc.add_paragraph()
                p.style = 'List Number'
                code_text = '\n'.join(code_lines)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 粗体和斜体处理的普通段落
        elif line.strip():
            # 处理内联格式
            paragraph = doc.add_paragraph()
            add_formatted_text(paragraph, line)
        
        i += 1


def add_formatted_text(paragraph, text):
    """
    添加带有格式的文本（粗体、斜体等）
    
    Args:
        paragraph: Word 段落对象
        text: 包含格式的文本
    """
    # 处理 **粗体** 和 *斜体*
    pattern = r'(\*\*.*?\*\*|\*.*?\*|\[.*?\]\(.*?\))'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        # 粗体
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        
        # 斜体
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        
        # 链接
        elif part.startswith('[') and '](' in part:
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                text_part = match.group(1)
                url_part = match.group(2)
                # Word 中添加超链接
                run = paragraph.add_run(text_part)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
            else:
                paragraph.add_run(part)
        
        # 普通文本
        else:
            paragraph.add_run(part)


def add_table_from_markdown(doc, table_lines):
    """
    从 Markdown 表格行创建 Word 表格
    
    Args:
        doc: Document 对象
        table_lines: 表格行列表
    """
    # 解析第一行（表头）
    header_line = table_lines[0].strip('|').split('|')
    header = [cell.strip() for cell in header_line]
    
    # 跳过分隔符行（第二行）
    # 获取数据行（第三行及以后）
    data_lines = table_lines[2:]
    
    # 创建表格
    num_cols = len(header)
    num_rows = len(data_lines) + 1  # +1 表头
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # 添加表头
    header_cells = table.rows[0].cells
    for i, cell_text in enumerate(header):
        header_cells[i].text = cell_text
        # 设置表头样式
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置表头背景色
        set_cell_background(header_cells[i], "0033CC")
    
    # 添加数据行
    for row_idx, data_line in enumerate(data_lines, start=1):
        row_data = data_line.strip('|').split('|')
        row_cells = table.rows[row_idx].cells
        
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                row_cells[col_idx].text = cell_text.strip()
                # 设置单元格对齐
                for paragraph in row_cells[col_idx].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_heading_color(heading, color):
    """
    设置标题颜色
    
    Args:
        heading: 标题段落
        color: RGBColor 对象
    """
    for run in heading.runs:
        run.font.color.rgb = color


def set_cell_background(cell, color):
    """
    设置表格单元格背景色
    
    Args:
        cell: 表格单元格
        color: 16 进制颜色代码（不含 #）
    """
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def extract_image_path(line):
    """
    从 Markdown 图片语法中提取路径
    
    Args:
        line: Markdown 图片语法行
    
    Returns:
        图片路径或 None
    """
    # ![alt](path) -> path
    match = re.match(r'!\[.*?\]\((.*?)\)', line)
    if match:
        return match.group(1)
    return None
